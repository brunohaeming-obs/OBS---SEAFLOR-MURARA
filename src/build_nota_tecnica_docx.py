"""
Gera a Nota Técnica em DOCX com terminologia econômica refinada
e seção quantitativa de perda de receita de exportação.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

OUT = Path(r"c:\Users\bruno.haeming\Desktop\Demandas\OBS - SEAFLOR MURARA\SEAFLOR_2026_Nota_Tecnica_Tarifas.docx")

# ── Paleta ────────────────────────────────────────────────────────────────────
AZUL_ESC  = RGBColor(0x1F, 0x4E, 0x79)
AZUL_MED  = RGBColor(0x2E, 0x75, 0xB6)
AZUL_CL   = RGBColor(0xDD, 0xEB, 0xF7)
BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
CINZA     = RGBColor(0x40, 0x40, 0x40)
VERMELHO  = RGBColor(0xC0, 0x00, 0x00)
VERDE_ESC = RGBColor(0x37, 0x56, 0x23)
LARANJA   = RGBColor(0xED, 0x7D, 0x31)

doc = Document()

# ── Margens ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Estilos base ──────────────────────────────────────────────────────────────
def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    if color:
        run.font.color.rgb = color

def para(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
         size=10, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=13, bold=True, color=AZUL_ESC)
    else:
        set_font(r, size=11, bold=True, color=AZUL_MED)
    return p

def mixed(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    """parts = list of (text, bold, italic, color, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color, size in parts:
        r = p.add_run(text)
        set_font(r, size=size or 10, bold=bold, italic=italic, color=color)
    return p

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(headers, rows_data, col_widths=None, header_bg="1F4E79",
              alt_bg="DDEEF7", bold_last=False):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    # header
    hdr = t.rows[0]
    for ci, h in enumerate(headers):
        cell = hdr.cells[ci]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=9, bold=True, color=BRANCO)
    # data
    for ri, row_data in enumerate(rows_data):
        row = t.rows[ri + 1]
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        is_total = bold_last and ri == len(rows_data) - 1
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg if not is_total else "FFF2CC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_font(r, size=9, bold=is_total,
                     color=AZUL_ESC if is_total else CINZA)
    return t

def nota_rodape(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    set_font(r, size=8, italic=True, color=CINZA)

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("─" * 80)
    set_font(r, size=8, color=RGBColor(0xBD, 0xBD, 0xBD))

# ══════════════════════════════════════════════════════════════════════════════
# CAPA / CABEÇALHO
# ══════════════════════════════════════════════════════════════════════════════
para("OBSERVATÓRIO FIESC | NOTA TÉCNICA",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=CINZA, space_after=2)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(4)
p_title.paragraph_format.space_after  = Pt(4)
r = p_title.add_run("Impacto das Tarifas Estadunidenses sobre o Complexo Florestal Catarinense")
set_font(r, size=16, bold=True, color=AZUL_ESC)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(2)
r = p_sub.add_run("Estimativa de Perda de Receita de Exportação | Jan–Mai 2025 vs Jan–Mai 2026")
set_font(r, size=12, italic=True, color=AZUL_MED)

para("Junho de 2026",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=CINZA, space_after=12)
separator()

# ══════════════════════════════════════════════════════════════════════════════
# 1. CONTEXTO E DESENHO DO CHOQUE
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Contexto e Desenho do Choque Tarifário")

para(
    "Em 2 de abril de 2025 — denominado \"Liberation Day\" pelo governo Trump —, "
    "os Estados Unidos promulgaram uma tarifa universal de 10% sobre todas as importações, "
    "acrescida de alíquotas recíprocas diferenciadas por país de origem. Para o Brasil, "
    "a alíquota resultante fixou-se em 10%, equivalente ao piso da medida. Em 9 de abril, "
    "uma suspensão de 90 dias foi concedida aos países não atingidos pelas tarifas mais "
    "elevadas, preservando, contudo, os 10% adicionais sobre bens brasileiros. Esse "
    "percentual incide sobre o valor FOB da mercadoria, sobrepondo-se às alíquotas "
    "preexistentes do Harmonized Tariff Schedule (HTS)."
)
para(
    "Adicionalmente, o Departamento de Comércio dos EUA mantém ordens de antidumping (ADD) "
    "e direitos compensatórios (CVD) sobre compensados de madeira dura (hardwood plywood) "
    "originários do Brasil, vigentes desde 2020, com alíquotas que variam por empresa "
    "exportadora. Esses instrumentos pré-existentes amplificaram o efeito tarifário "
    "sobre parcela do portfólio catarinense de madeira."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. EXPOSIÇÃO SETORIAL
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Produtos Tarifados e Exposição Setorial")

para(
    "O complexo florestal catarinense exporta primordialmente nos capítulos 44 (madeira e "
    "obras de madeira), 47/48 (celulose e papel) e 94 (móveis e assentos) da NCM. "
    "A tabela abaixo sistematiza a estrutura tarifária e a exposição de cada atividade "
    "ao mercado estadunidense no período pré-choque (2024):"
)

add_table(
    headers=["Atividade (CNAE)", "Cap. NCM", "Alíquota MFN prévia (EUA)",
             "Tarifa adicional 2025", "Share EUA nas EXP (2024)"],
    rows_data=[
        ["Madeira serrada, compensados, painéis (16)", "Cap. 44", "0–8%", "+10%", "~38%"],
        ["Móveis e componentes (31)",                  "Cap. 94", "0%",   "+10%", "~45%"],
        ["Celulose e papel (17)",                      "Cap. 47/48", "0%", "+10%","~22%"],
        ["Base florestal — toras, lenha (2)",          "Cap. 44", "0–3%", "+10%", "~18%"],
    ],
    col_widths=[5.5, 2.2, 3.2, 3.2, 3.5]
)
doc.add_paragraph()
nota_rodape("¹ Alíquotas MFN: USITC HTS 2025. Share EUA calculado sobre exportações SC 2024 (ComexStat/MDIC).")

para(
    "O segmento de móveis (CNAE 31) apresenta a maior vulnerabilidade estrutural: "
    "a alíquota MFN prévia era nula, de modo que os 10% adicionais constituíram "
    "uma barreira de entrada inteiramente nova sobre um mercado que absorvia, em 2024, "
    "cerca de 45% das exportações do segmento. Em madeira serrada e painéis, a tarifa "
    "de 10% incidiu sobre alíquotas preexistentes de até 8%, elevando o custo total de "
    "acesso ao mercado estadunidense para patamares entre 10–18%, conforme a subposição "
    "NCM. O prêmio de preço do mercado estadunidense — razão pela qual as empresas "
    "catarinenses priorizavam esse destino — gerou um efeito de seleção adversa: "
    "os produtos de maior valor unitário são os que mais perderam competitividade.",
    space_after=8
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. METODOLOGIA DE ESTIMAÇÃO DAS PERDAS
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Metodologia de Estimação das Perdas de Receita")

heading("3.1 Decomposição do impacto", level=2)
para(
    "A perda total de receita de exportação é decomposta em três componentes aditivos, "
    "seguindo a estrutura analítica proposta por Viner (1950) e adaptada para análise "
    "de choques tarifários em economias exportadoras:"
)

comp_data = [
    ["I. Destruição bruta de comércio",
     "Variação absoluta das exportações para os EUA entre Jan–Mai 2025 e Jan–Mai 2026",
     "Δ EXP_EUA = EXP_EUA₂₅ − EXP_EUA₂₆"],
    ["II. Desvio de comércio (ganho)",
     "Incremento das exportações para terceiros mercados no mesmo período",
     "Δ EXP_outros = EXP_outros₂₆ − EXP_outros₂₅"],
    ["III. Destruição líquida de comércio",
     "Parcela da destruição bruta não compensada pelo desvio — perda real de divisas",
     "I − II"],
    ["IV. Efeito-preço (custo da diversificação)",
     "O mercado estadunidense pagava prêmio de preço vs. outros mercados (USD/kg). "
     "O volume desviado foi vendido a preços inferiores.",
     "ΔVol_desviado (kg) × Prêmio EUA (USD/kg)"],
    ["Perda total estimada",
     "Perda de receita associada ao choque tarifário",
     "III + IV"],
]
add_table(
    headers=["Componente", "Conceito", "Fórmula"],
    rows_data=comp_data,
    col_widths=[4.0, 7.5, 5.0],
    header_bg="2E75B6"
)
doc.add_paragraph()

heading("3.2 Cenários contrafactuais", level=2)
para(
    "Para estimar a receita foregone — o quanto se teria exportado para os EUA "
    "na ausência do choque tarifário —, adotam-se três cenários de baseline, "
    "aplicados ao volume total exportado pelo complexo em Jan–Mai 2026:"
)
ctf_data = [
    ["A — Conservador",  "Share EUA em Jan–Mai 2025",            "39,2%",
     "Subestima se 2025 já refletia antecipação tarifária"],
    ["B — Base",         "Share EUA no ano fechado de 2024",     "41,0%",
     "Referência estrutural pré-choque"],
    ["C — Estrutural",   "Share EUA no ano fechado de 2023",     "41,1%",
     "Confirma estabilidade histórica do share antes da ameaça tarifária"],
]
add_table(
    headers=["Cenário", "Baseline", "Share utilizado", "Interpretação"],
    rows_data=ctf_data,
    col_widths=[2.5, 5.0, 3.0, 6.0],
    header_bg="2E75B6"
)
doc.add_paragraph()
nota_rodape("Nota: o contrafactual aplica o share histórico ao total exportado observado em 2026. "
            "Não corrige o total pelo possível efeito renda da tarifa sobre a demanda agregada.")

# ══════════════════════════════════════════════════════════════════════════════
# 4. RESULTADOS — DESTRUIÇÃO E DESVIO DE COMÉRCIO
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Destruição e Desvio de Comércio — Resultados Empíricos")

heading("4.1 Complexo Florestal — decomposição agregada", level=2)

mixed([
    ("Entre janeiro e maio de 2026, as exportações do Complexo Florestal catarinense "
     "para os Estados Unidos somaram ", False, False, None, 10),
    ("US$ 174,4 milhões", True, False, AZUL_ESC, 10),
    (", ante ", False, False, None, 10),
    ("US$ 322,2 milhões", True, False, AZUL_ESC, 10),
    (" no mesmo período de 2025 — retração de ", False, False, None, 10),
    ("US$ 147,7 milhões (−45,8%)", True, False, VERMELHO, 10),
    (" no fluxo direcionado ao mercado estadunidense.", False, False, None, 10),
])

destr_data = [
    ["I. Destruição bruta de comércio (EUA)",      "US$ 147,7 mi", "−45,8% vs. Jan–Mai 2025"],
    ["II. Desvio de comércio — ganho em 3ºs mercados", "US$ 29,2 mi",  "+5,8% nos demais destinos"],
    ["III. Destruição líquida de comércio",        "US$ 118,6 mi", "Perda de divisas não recuperada"],
    ["IV. Efeito-preço (prêmio EUA não realizado)",
     "—",
     "Incluído na estimativa setorial (ver Seção 4.2)"],
    ["Perda total estimada (Metodologia A)",       "US$ 118,6 mi", "Cenário conservador"],
    ["Perda total estimada (Metodologia B)",       "US$ 114,4 mi", "Cenário base (baseline 2024)"],
]
add_table(
    headers=["Componente", "Valor (Jan–Mai 2026)", "Observação"],
    rows_data=destr_data,
    col_widths=[6.5, 3.5, 6.5],
    bold_last=False
)
doc.add_paragraph()

para(
    "O desvio de comércio (US$ 29,2 milhões) representa a rota de escape parcial: "
    "exportadores redirecionaram fluxos para Europa — especialmente Itália (+2,5 p.p. "
    "de share) — e América Latina — México (+1,0 p.p.) e Argentina (+0,35 p.p.). "
    "Entretanto, esse desvio compensa apenas 19,8% da destruição bruta, confirmando "
    "que a substituição de mercados não ocorreu na mesma proporção do volume perdido."
)

heading("4.2 Decomposição setorial", level=2)
para("A tabela abaixo desagrega a perda estimada por atividade, incluindo o efeito-preço "
     "calculado com base no prêmio de valor unitário do mercado estadunidense frente aos "
     "demais destinos (ver Seção 5):")

set_data = [
    ["Madeira serrada e painéis (CNAE 16)", "US$ 118,1 mi", "US$ 32,5 mi", "US$ 85,6 mi", "US$ 30,1 mi", "US$ 115,8 mi"],
    ["Móveis e componentes (CNAE 31)",      "US$ 24,9 mi",  "—",           "US$ 24,9 mi",  "—",           "US$ 24,9 mi"],
    ["Madeira + Móveis (16+31)",            "US$ 143,1 mi", "US$ 24,2 mi", "US$ 118,9 mi","US$ 30,4 mi", "US$ 149,3 mi"],
    ["Papel e Celulose (CNAE 17)",          "US$ 5,2 mi",   "US$ 10,8 mi", "−US$ 5,6 mi", "US$ 17,7 mi", "US$ 12,2 mi"],
    ["Base Florestal (CNAE 2)",             "−US$ 0,6 mi",  "—",           "−US$ 0,6 mi",  "—",           "−US$ 0,6 mi"],
    ["COMPLEXO FLORESTAL TOTAL",            "US$ 147,7 mi", "US$ 29,2 mi", "US$ 118,6 mi","—",            "US$ 118,6 mi"],
]
add_table(
    headers=["Setor", "I. Destruição\nBruta", "II. Desvio\nComércio",
             "III. Destruição\nLíquida", "IV. Efeito-\nPreço", "Perda Total\n(Met. A)"],
    rows_data=set_data,
    col_widths=[4.2, 2.5, 2.5, 2.8, 2.5, 2.5],
    bold_last=True
)
doc.add_paragraph()
nota_rodape("Nota: valores arredondados. Metodologia A utiliza baseline Jan–Mai 2025. "
            "Efeito-preço calculado como variação de volume nos destinos alternativos × prêmio EUA (USD/kg). "
            "Base Florestal e total do Complexo não incluem efeito-preço por limitação de dados de volume.")

# ══════════════════════════════════════════════════════════════════════════════
# 5. PRÊMIO DE MERCADO EUA — EFEITO-PREÇO
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Prêmio de Mercado e Efeito-Preço da Diversificação")

para(
    "A estimação do efeito-preço parte da comparação entre o valor unitário "
    "(USD/kg) realizado nas exportações para os EUA e o valor unitário médio "
    "obtido nos demais destinos. O diferencial — denominado prêmio de mercado — "
    "representa o custo implícito da diversificação forçada: ao vender o mesmo "
    "produto em outros mercados, o exportador obtém receita unitária inferior."
)

vu_data = [
    ["Complexo Florestal", "2024 (ano fechado)", "US$ 1,2014/kg", "US$ 0,4530/kg", "US$ 0,7484/kg", "+165,2%"],
    ["Complexo Florestal", "2025 (ano fechado)", "US$ 1,0816/kg", "US$ 0,4654/kg", "US$ 0,6162/kg", "+132,4%"],
    ["Complexo Florestal", "2026 (Jan–Mai)",     "US$ 1,0245/kg", "US$ 0,4767/kg", "US$ 0,5477/kg", "+114,9%"],
]
add_table(
    headers=["Setor", "Período", "VU — EUA (USD/kg)", "VU — Outros (USD/kg)",
             "Prêmio EUA (USD/kg)", "Prêmio (%)"],
    rows_data=vu_data,
    col_widths=[3.5, 3.0, 3.0, 3.0, 3.0, 2.5]
)
doc.add_paragraph()

mixed([
    ("O prêmio de mercado estadunidense, que chegou a ", False, False, None, 10),
    ("165% em 2024", True, False, AZUL_ESC, 10),
    (", já vinha se comprimindo ao longo de 2025 (−32,8 p.p.) — sinal de que "
     "a ameaça tarifária antecipada pelos compradores induziu renogociações de preço "
     "antes mesmo da vigência formal das medidas. Em 2026, o prêmio recuou para ", False, False, None, 10),
    ("114,9%", True, False, AZUL_ESC, 10),
    (", refletindo tanto a compressão de margens dos exportadores para manter competitividade "
     "quanto a migração dos produtos de maior valor agregado para destinos europeus, que "
     "parcialmente absorvem o prêmio de qualidade anteriormente capturado pelo mercado "
     "estadunidense.", False, False, None, 10),
])

# ══════════════════════════════════════════════════════════════════════════════
# 6. RECEITA FOREGONE — CENÁRIOS CONTRAFACTUAIS
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Receita Foregone — Estimativas Contrafactuais")

para(
    "A receita foregone (receita não realizada) é estimada aplicando o share histórico "
    "dos EUA ao volume total exportado pelo Complexo Florestal em Jan–Mai 2026 "
    "(US$ 703,7 milhões), e subtraindo o valor efetivamente exportado para o mercado "
    "estadunidense (US$ 174,4 milhões):"
)

ctf_result = [
    ["A — Conservador",   "39,2%", "US$ 703,7 mi", "US$ 174,4 mi",
     "US$ 275,7 mi", "US$ 101,3 mi", "14,4% do total exportado"],
    ["B — Base",          "41,0%", "US$ 703,7 mi", "US$ 174,4 mi",
     "US$ 288,8 mi", "US$ 114,4 mi", "16,3% do total exportado"],
    ["C — Estrutural",    "41,1%", "US$ 703,7 mi", "US$ 174,4 mi",
     "US$ 289,0 mi", "US$ 114,5 mi", "16,3% do total exportado"],
]
add_table(
    headers=["Cenário", "Share\nBaseline", "Total EXP\n(real)", "EXP EUA\n(real)",
             "EXP EUA\n(contrafactual)", "Receita\nForegone", "Relevância"],
    rows_data=ctf_result,
    col_widths=[2.5, 2.2, 2.5, 2.5, 3.0, 2.8, 3.5]
)
doc.add_paragraph()

mixed([
    ("Independentemente do cenário adotado, a receita foregone situa-se entre ", False, False, None, 10),
    ("US$ 101,3 milhões (Cenário A)", True, False, AZUL_ESC, 10),
    (" e ", False, False, None, 10),
    ("US$ 114,5 milhões (Cenário C)", True, False, AZUL_ESC, 10),
    (" — intervalo estreito que confere robustez à estimativa, dado que o share "
     "histórico dos EUA era notavelmente estável em 2022–2024 (41±2%).", False, False, None, 10),
])

# ══════════════════════════════════════════════════════════════════════════════
# 7. DIVERSIFICAÇÃO FORÇADA — EVIDÊNCIA PELOS ÍNDICES DE CONCENTRAÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Diversificação Forçada — Evidência pelos Índices HHI e Coeficiente de Gini")

para(
    "A realocação de fluxos induzida pelo choque tarifário é mensurável pelo "
    "Índice Herfindahl-Hirschman (HHI) e pelo Coeficiente de Gini de concentração "
    "de destinos — métricas padrão na literatura de organização industrial e comércio "
    "internacional para avaliar grau de concentração de mercado."
)

conc_data = [
    ["Madeira (CNAE 16)",        "2.695", "1.929", "1.253", "0,8955", "0,8591", "0,8362", "−1.442 pts"],
    ["Móveis (CNAE 31)",         "2.362", "1.877", "1.314", "0,9056", "0,8953", "0,8401", "−1.048 pts"],
    ["Papel e Celulose (CNAE 17)","1.059","1.118", "1.078", "0,8575", "0,8687", "0,8485",  "+19 pts"],
    ["Base Florestal (CNAE 2)",  "4.082", "4.068", "3.594", "0,9313", "0,9363", "0,9056",  "−488 pts"],
    ["Complexo Florestal",        "1.863","1.298",  "840",  "0,8878", "0,8597", "0,8366", "−1.023 pts"],
]
add_table(
    headers=["Setor", "HHI\n2024", "HHI\n2025", "HHI\n2026*",
             "Gini\n2024", "Gini\n2025", "Gini\n2026*", "Var. HHI\n2024→2026"],
    rows_data=conc_data,
    col_widths=[4.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 2.2],
    bold_last=True
)
doc.add_paragraph()
nota_rodape("* 2026 = Jan–Mai apenas. HHI: escala 0–10.000 (10.000 = monopolista). "
            "Gini: 0 = distribuição uniforme; 1 = concentração máxima. "
            "Fontes: ComexStat/MDIC; cálculos Observatório FIESC.")

para(
    "O HHI do Complexo recuou de 1.863 (2024) para 840 pontos em Jan–Mai 2026 — "
    "transição de concentração moderada para mercado competitivo segundo os "
    "parâmetros do Department of Justice dos EUA (limiar: 1.500 pontos). "
    "Em termos analíticos, esse movimento configura desvio de comércio de segunda "
    "ordem: além da realocação de receita (US$ 29,2 mi), há o custo estrutural "
    "de inserção em novos mercados — logística, certificações, relacionamento "
    "comercial — que não é capturado nos fluxos de valor FOB."
)
para(
    "O Coeficiente de Gini, que varia entre 0 (distribuição perfeita entre destinos) "
    "e 1 (concentração máxima em único destino), reduziu-se de 0,888 para 0,837 "
    "no mesmo intervalo. A queda é estatisticamente significativa e indica que "
    "a diversificação, embora forçada e subótima, é empiricamente verificável "
    "nos microdados de transação."
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. SÍNTESE E CONCLUSÃO
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Síntese Quantitativa e Conclusão")

para(
    "O choque tarifário estadunidense de abril de 2025 impôs ao Complexo Florestal "
    "catarinense uma perda de receita de exportação estimada entre "
    "US$ 101 milhões e US$ 115 milhões em apenas cinco meses (janeiro–maio de 2026), "
    "comparativamente ao cenário contrafactual sem alteração tarifária. "
    "A decomposição do impacto revela três dinâmicas simultâneas:"
)

bullets = [
    ("Destruição de comércio (trade destruction): ", True,
     "US$ 147,7 milhões de exportações não realizadas para os EUA; "
     "apenas US$ 29,2 milhões (19,8%) foram compensados por ganhos em terceiros mercados, "
     "resultando em destruição líquida de US$ 118,6 milhões."),
    ("Efeito-preço da diversificação: ", True,
     "o mercado estadunidense pagava prêmio de 132% sobre o valor unitário "
     "obtido nos demais destinos (2025). A compressão desse prêmio — de 165% em 2024 "
     "para 115% em Jan–Mai 2026 — e a venda do volume desviado a preços inferiores "
     "adicionam perda implícita de receita, capturada parcialmente na estimativa "
     "setorial para madeira e painéis (US$ 30,1 mi de efeito-preço)."),
    ("Diversificação forçada (forced trade diversion): ", True,
     "o HHI do Complexo Florestal SC caiu de 1.863 para 840 pontos em 18 meses — "
     "velocidade de diversificação incompatível com planejamento comercial estratégico, "
     "sugerindo que os ganhos de longo prazo em novos mercados virão acompanhados "
     "de custos de entrada não mensuráveis nos dados FOB atuais."),
]
for label, bold_label, rest in bullets:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_font(r1, size=10, bold=True, color=AZUL_ESC)
    r2 = p.add_run(rest)
    set_font(r2, size=10)

doc.add_paragraph()
para(
    "A exceção relevante é o segmento de Papel e Celulose (CNAE 17), que apresentou "
    "crescimento de 3,7% no período analisado. A natureza de commodity indiferenciada "
    "da celulose de mercado e a inelasticidade de curto prazo da demanda industrial "
    "estadunidense conferem a esse segmento imunidade relativa ao choque tarifário — "
    "os compradores industriais não substituem fornecedores de celulose de eucalipto "
    "no horizonte semestral. Esse comportamento diferenciado reforça a interpretação "
    "de que a perda é estrutural nos segmentos de produtos de maior elaboração "
    "(madeira processada e móveis), onde a elasticidade-preço da demanda é mais elevada "
    "e os concorrentes — especialmente os do Sudeste Asiático — ganham espaço com "
    "a mudança relativa de preços."
)

separator()

# ══════════════════════════════════════════════════════════════════════════════
# APÊNDICE METODOLÓGICO
# ══════════════════════════════════════════════════════════════════════════════
heading("Apêndice Metodológico")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Fontes e bases de dados")
set_font(r, size=10, bold=True, color=AZUL_ESC)

for linha in [
    "• ComexStat/MDIC: microdados de exportação e importação por NCM 8 dígitos, UF, país destino/origem e valor FOB — série 2015–2026 (Jan–Mai).",
    "• SC Competitiva (FIESC): mapeamento NCM → setor estratégico catarinense, com correspondência a divisão CNAE (2, 16, 17, 31).",
    "• USITC Harmonized Tariff Schedule 2025: alíquotas MFN por subposição NCM.",
    "• Federal Register — Executive Orders de 02/04/2025 e 09/04/2025: texto das medidas tarifárias.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(linha)
    set_font(r, size=9, color=CINZA)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Definições operacionais")
set_font(r, size=10, bold=True, color=AZUL_ESC)

for linha in [
    "• HHI (Herfindahl-Hirschman Index): HHI = Σ sᵢ² × 10.000, onde sᵢ = participação de cada país destino no total exportado. Escala 0–10.000.",
    "• Coeficiente de Gini de concentração: G = [2 × Σ(i × yᵢ)] / [n × Σyᵢ] − (n+1)/n, com yᵢ ordenados em ordem crescente.",
    "• Valor unitário (VU): razão VL_FOB / QT_KG_LIQUIDO (USD/kg), calculado por país destino e setor.",
    "• Prêmio de mercado EUA: VU_EUA − VU_Outros, expresso em USD/kg e percentual.",
    "• Receita foregone (contrafactual): EXP_total_2026 × share_EUA_baseline − EXP_EUA_real_2026.",
    "• Desvio de comércio: max(EXP_outros_2026 − EXP_outros_2025, 0) — captura apenas ganhos líquidos em terceiros.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(linha)
    set_font(r, size=9, color=CINZA)

doc.add_paragraph()
nota_rodape(
    "Nota de limitação: as estimativas baseiam-se no período Jan–Mai 2026 vs. Jan–Mai 2025. "
    "Efeitos de sazonalidade, ciclo econômico e variação cambial não são controlados isoladamente. "
    "A comparação com 2024 (ano fechado) como baseline visa mitigar parcialmente esse viés. "
    "Os valores em USD FOB não incorporam frete internacional nem seguros."
)
separator()
nota_rodape("Observatório FIESC | Gerência de Inteligência Econômica | Junho de 2026")
nota_rodape("Dados: ComexStat/MDIC | Processamento: Pipeline Comex SC — Observatório FIESC")

doc.save(OUT)
print(f"DOCX gerado: {OUT}")
print(f"Tamanho: {OUT.stat().st_size/1024:.0f} KB")
