"""
Gera o roteiro da apresentação SEAFLOR 2026 em formato .docx.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(r"c:\Users\bruno.haeming\Desktop\Demandas\OBS - SEAFLOR MURARA\SEAFLOR_2026_Roteiro_Apresentacao.docx")

AZUL_FIESC   = RGBColor(0x1F, 0x4E, 0x79)
AZUL_CLARO   = RGBColor(0x2E, 0x75, 0xB6)
CINZA_TEXTO  = RGBColor(0x40, 0x40, 0x40)
VERMELHO     = RGBColor(0xC0, 0x00, 0x00)
VERDE        = RGBColor(0x37, 0x56, 0x23)
BRANCO       = RGBColor(0xFF, 0xFF, 0xFF)


# ── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, italic=False,
            color=None, size=None, highlight=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = AZUL_FIESC if level == 1 else AZUL_CLARO
        run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(doc, text, bold=False, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, text, bold=bold, italic=italic,
            color=color or CINZA_TEXTO, size=10.5)
    return p


def quote_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'"{text}"')
    run.bold             = True
    run.italic           = True
    run.font.color.rgb   = AZUL_FIESC
    run.font.size        = Pt(11)


def add_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], "1F4E79")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold           = True
            run.font.color.rgb = BRANCO
            run.font.size      = Pt(9.5)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linhas
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg    = "DEEAF1" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            para = cells[ci].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9.5)
                if str(val).startswith("-") or str(val).startswith("↓"):
                    run.font.color.rgb = VERMELHO
                elif str(val).startswith("+") or str(val).startswith("↑"):
                    run.font.color.rgb = VERDE

    # Larguras
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def slide_title(doc, number, title, subtitle=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"SLIDE {number}")
    run.bold           = True
    run.font.color.rgb = AZUL_CLARO
    run.font.size      = Pt(9)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    r = p2.add_run(title)
    r.bold           = True
    r.font.color.rgb = AZUL_FIESC
    r.font.size      = Pt(13)

    if subtitle:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(4)
        r3 = p3.add_run(subtitle)
        r3.italic        = True
        r3.font.color.rgb = CINZA_TEXTO
        r3.font.size     = Pt(10)


def label(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(txt)
    r.bold           = True
    r.font.color.rgb = AZUL_CLARO
    r.font.size      = Pt(10)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size      = Pt(10)
    r.font.color.rgb = CINZA_TEXTO


def divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = AZUL_CLARO
        run.font.size      = Pt(7)


def bloco(doc, texto):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(texto)
    run.bold           = True
    run.font.size      = Pt(13)
    run.font.color.rgb = BRANCO
    # sombreamento simulado com cor de fonte (word não tem fundo fácil em parágrafo)
    # usamos heading 1 com cor diferente
    doc.paragraphs[-1].clear()
    h = doc.add_heading(texto, level=1)
    for r in h.runs:
        r.font.color.rgb = AZUL_FIESC
        r.font.size      = Pt(14)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after  = Pt(4)


# ── documento ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Margens
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Capa ──────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(40)
    r = p_title.add_run("SEAFLOR 2026")
    r.bold           = True
    r.font.size      = Pt(28)
    r.font.color.rgb = AZUL_FIESC

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Roteiro da Apresentação")
    r2.font.size      = Pt(16)
    r2.font.color.rgb = AZUL_CLARO

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Comércio Exterior SC — Complexo Florestal")
    r3.font.size      = Pt(13)
    r3.font.color.rgb = CINZA_TEXTO

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, sep in [("Semana de Aperfeiçoamento em Engenharia Florestal", " | "),
                     ("UFPR", " | "), ("24–26 jun 2026", "")]:
        r4 = p4.add_run(txt + sep)
        r4.font.size      = Pt(10)
        r4.font.color.rgb = CINZA_TEXTO

    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("Representação: FIESC e ACR")
    r5.bold           = True
    r5.font.color.rgb = AZUL_FIESC
    r5.font.size      = Pt(11)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r6 = p6.add_run("Fonte: ComexStat/MDIC | Observatório FIESC | Junho 2026")
    r6.italic        = True
    r6.font.size     = Pt(9)
    r6.font.color.rgb = CINZA_TEXTO

    doc.add_page_break()

    # ── Fio condutor ──────────────────────────────────────────────────────────
    heading(doc, "Fio Condutor Narrativo", level=1)
    quote_box(doc, "SC tem uma das bases florestais mais produtivas do mundo. "
                   "O que fazemos com isso no mercado internacional — e o que nos freia.")
    body(doc, "A apresentação percorre três movimentos:", space_after=2)
    bullet(doc, "O que temos — base produtiva e expressão exportadora")
    bullet(doc, "Para onde vai — destinos, concentração e vulnerabilidades")
    bullet(doc, "O que nos ameaça — tarifas, desvio de comércio e agenda de diversificação")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCO 1
    # ══════════════════════════════════════════════════════════════════════════
    bloco(doc, "BLOCO 1 — Contexto e Importância do Setor")

    # Slide 1
    divider(doc)
    slide_title(doc, 1, "SC Florestal: US$ 1,86 bilhão exportados em 2025")
    label(doc, "Mensagem central")
    body(doc, "O complexo florestal plantado respondeu por 15,3% de todas as exportações "
              "de Santa Catarina em 2025 — mais de US$ 1,8 bilhão em produtos de alto "
              "valor agregado (madeira, móveis, celulose e produção florestal).")
    label(doc, "Visual sugerido")
    body(doc, "Número grande centralizado + mapa do Brasil destacando SC")
    label(doc, "Aba de referência")
    body(doc, "5_Participacao_SC")

    # Slide 2
    divider(doc)
    slide_title(doc, 2, "Três segmentos, um ativo em comum: a floresta plantada")
    label(doc, "Tabela de impacto")
    add_table(doc,
        ["Segmento", "CNAE", "Exportações 2025"],
        [
            ["Madeira e produtos de madeira", "16", "US$ 1,17 bi"],
            ["Móveis",                         "31", "US$ 253 mi"],
            ["Papel e Celulose",               "17", "US$ 371 mi"],
            ["Produção Florestal (base)",       "2",  "US$ 49 mi"],
            ["TOTAL COMPLEXO",                 "—",  "US$ 1,86 bi"],
        ],
        col_widths=[8, 2.5, 5],
    )
    label(doc, "Fala do orador")
    quote_box(doc, "Quando falamos em floresta plantada, não estamos falando só de toras. "
                   "Estamos falando de móveis nos EUA, celulose na Argentina, compensados na Europa.")
    label(doc, "Aba de referência")
    body(doc, "1_Serie_Historica")

    # Slide 3
    divider(doc)
    slide_title(doc, 3, "Série histórica (2015–2025)",
                "Uma década de crescimento — com turbulências")
    label(doc, "Narrativa por período")
    bullet(doc, "2015–2019: crescimento gradual, US$ 1,06 bi → US$ 1,43 bi")
    bullet(doc, "2020–2022: boom pós-COVID (construção civil americana); pico em 2022: US$ 2,31 bi")
    bullet(doc, "2023: acomodação (-24%); mercado americano desacelera")
    bullet(doc, "2024–2025: estabilização em torno de US$ 1,86–1,94 bi")
    label(doc, "Visual sugerido")
    body(doc, "Gráfico de linha com três séries (Madeira, Móveis, Papel) + barra de saldo no fundo")
    label(doc, "Ponto de atenção")
    body(doc, "O pico de 2021-2022 foi excepcional. O nível atual ainda é 75% acima do pré-pandemia.")
    label(doc, "Aba de referência")
    body(doc, "1_Serie_Historica")

    # Slide 4
    divider(doc)
    slide_title(doc, 4, "Saldo comercial superavitário consistente",
                "O complexo florestal SC gera mais de US$ 1,3 bi de saldo positivo todo ano")
    add_table(doc,
        ["Ano", "Exportações", "Importações", "Saldo"],
        [
            ["2019", "US$ 1,43 bi", "US$ 264 mi", "US$ 1,16 bi"],
            ["2021", "US$ 2,15 bi", "US$ 382 mi", "US$ 1,76 bi"],
            ["2022", "US$ 2,31 bi", "US$ 322 mi", "US$ 1,99 bi"],
            ["2024", "US$ 1,94 bi", "US$ 434 mi", "US$ 1,51 bi"],
            ["2025", "US$ 1,86 bi", "US$ 478 mi", "US$ 1,38 bi"],
        ],
        col_widths=[2.5, 4, 4, 4],
    )
    label(doc, "Alerta")
    body(doc, "Importações cresceram +10% de 2024 para 2025, enquanto exportações caíram 4,5%. "
              "Saldo ainda positivo, mas a tendência merece monitoramento.", color=VERMELHO)
    label(doc, "Aba de referência")
    body(doc, "2_Saldo_Comercial")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCO 2
    # ══════════════════════════════════════════════════════════════════════════
    bloco(doc, "BLOCO 2 — Para onde vai: Destinos e Concentração")

    # Slide 5
    divider(doc)
    slide_title(doc, 5, "Top produtos exportados (2025)",
                "Madeira serrada lidera — mas o portfólio é diversificado")
    add_table(doc,
        ["Produto", "US$ FOB 2025", "CNAE"],
        [
            ["Madeira serrada",               "US$ 368 mi", "16"],
            ["Carpintaria para construção",   "US$ 254 mi", "16"],
            ["Madeira compensada",            "US$ 240 mi", "16"],
            ["Outros móveis",                 "US$ 240 mi", "31"],
            ["Madeira em forma",              "US$ 84 mi",  "16"],
            ["MDF",                           "US$ 81 mi",  "16"],
            ["Lenha e resíduos de madeira",   "US$ 44 mi",  "16"],
            ["Painéis aglomerados",           "US$ 38 mi",  "16"],
            ["Folheados",                     "US$ 26 mi",  "16"],
        ],
        col_widths=[9, 4, 2.5],
    )
    label(doc, "Fala do orador")
    quote_box(doc, "O produto mais exportado — madeira serrada — vai principalmente para o mercado "
                   "americano de construção civil. É aí que mora o risco.")
    label(doc, "Aba de referência")
    body(doc, "3_Top_Produtos_2025")

    # Slide 6
    divider(doc)
    slide_title(doc, 6, "Destinos por segmento (2025)",
                "Três segmentos, três perfis de mercado completamente diferentes")
    label(doc, "Madeira (CNAE 16) — concentrada em EUA e Europa")
    add_table(doc,
        ["País", "US$ FOB"],
        [["Estados Unidos", "US$ 596 mi (51%)"], ["México", "US$ 98 mi"],
         ["Reino Unido", "US$ 58 mi"], ["Itália", "US$ 48 mi"], ["China", "US$ 45 mi"]],
        col_widths=[8, 5],
    )
    label(doc, "Móveis (CNAE 31) — EUA + mercados emergentes")
    add_table(doc,
        ["País", "US$ FOB"],
        [["Estados Unidos", "US$ 101 mi (40%)"], ["Argentina", "US$ 30 mi"],
         ["Reino Unido", "US$ 22 mi"], ["Austrália", "US$ 15 mi"]],
        col_widths=[8, 5],
    )
    label(doc, "Papel e Celulose (CNAE 17) — América Latina")
    add_table(doc,
        ["País", "US$ FOB"],
        [["Argentina", "US$ 102 mi (27%)"], ["Paraguai", "US$ 38 mi"],
         ["Chile", "US$ 29 mi"], ["México", "US$ 24 mi"]],
        col_widths=[8, 5],
    )
    label(doc, "Base Florestal (CNAE 2) — Portugal e China")
    add_table(doc,
        ["País", "US$ FOB"],
        [["Portugal", "US$ 23 mi (47%)"], ["China", "US$ 21 mi (43%)"]],
        col_widths=[8, 5],
    )
    label(doc, "Visual sugerido")
    body(doc, "4 gráficos de barras horizontais em grade 2×2")
    label(doc, "Abas de referência")
    body(doc, "7_Destinos_Categoria | 4_Top_Destinos_2025")

    # Slide 7
    divider(doc)
    slide_title(doc, 7, "Coeficiente de Gini: concentração de destinos",
                "Quanto mais concentrado, mais vulnerável — e todos os segmentos estão concentrados")
    label(doc, "O que é o Gini aqui")
    body(doc, "Mede a desigualdade na distribuição das exportações entre os países compradores.")
    bullet(doc, "Gini = 0: todos os países compram o mesmo valor (máxima diversificação)")
    bullet(doc, "Gini = 1: um único país compra tudo (máxima concentração)")
    label(doc, "Resultados 2024 vs 2025")
    add_table(doc,
        ["Segmento", "Gini 2024", "Gini 2025", "Variação", "Nº Países"],
        [
            ["Base Florestal (CNAE 2)",       "0,9313", "0,9363", "+0,005 ↑", "42 → 47"],
            ["Móveis (CNAE 31)",              "0,9056", "0,8953", "-0,010 ↓", "106 → 105"],
            ["Madeira (CNAE 16)",             "0,8955", "0,8591", "-0,036 ↓", "131 → 126"],
            ["Papel e Celulose (CNAE 17)",    "0,8575", "0,8687", "+0,011 ↑", "105 → 105"],
        ],
        col_widths=[6.5, 2.5, 2.5, 2.5, 2.8],
    )
    label(doc, "Leituras-chave")
    bullet(doc, "Madeira e Móveis reduziram concentração em 2025 — efeito involuntário das tarifas americanas")
    bullet(doc, "Base Florestal: top 5 países = 96–97% das exportações (Portugal + China dominam)")
    bullet(doc, "Papel e Celulose é o mais diversificado — estruturalmente mais resiliente")
    bullet(doc, "Todos os segmentos ainda em faixa de alta concentração (Gini > 0,85)")
    label(doc, "Aba de referência")
    body(doc, "12_Gini_Resumo")

    # Slide 8
    divider(doc)
    slide_title(doc, 8, "HHI complementar ao Gini",
                "O HHI confirma: Madeira avançou na diversificação, mas ainda é vulnerável")
    label(doc, "Referências de interpretação do HHI")
    bullet(doc, "HHI > 2.500: mercado altamente concentrado")
    bullet(doc, "HHI 1.500–2.500: moderadamente concentrado")
    bullet(doc, "HHI < 1.500: desconcentrado")
    add_table(doc,
        ["Segmento", "HHI 2024", "HHI 2025", "Classificação 2025"],
        [
            ["Base Florestal (CNAE 2)",    "4.082", "4.068", "Alta concentração"],
            ["Madeira (CNAE 16)",          "2.695", "1.929", "Alta → Moderada"],
            ["Móveis (CNAE 31)",           "2.362", "1.877", "Alta → Moderada"],
            ["Papel e Celulose (CNAE 17)", "1.059", "1.118", "Desconcentrado"],
        ],
        col_widths=[6.5, 2.8, 2.8, 5],
    )
    label(doc, "Mensagem")
    quote_box(doc, "As tarifas americanas de 2025 forçaram uma diversificação que políticas de décadas "
                   "não conseguiram. É uma oportunidade disfarçada — se o setor souber aproveitar.")
    label(doc, "Aba de referência")
    body(doc, "12_Gini_Resumo")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCO 3
    # ══════════════════════════════════════════════════════════════════════════
    bloco(doc, "BLOCO 3 — O que nos ameaça: Tarifas e Agenda Futura")

    # Slide 9
    divider(doc)
    slide_title(doc, 9, "A queda abrupta do mercado americano",
                "O 'Liberation Day' de Trump chegou nas exportações de SC")
    label(doc, "Linha do tempo")
    bullet(doc, "02/abr/2025: Trump anuncia tarifas 'recíprocas' — Brasil recebe alíquota base de 10%")
    bullet(doc, "Jan–Jun/2025: EUA ainda representa ~40% das exportações florestais SC (normal)")
    bullet(doc, "Jul–Nov/2025: share cai para 18,5% em novembro — queda de quase 22 pp")
    label(doc, "Evolução mensal do share EUA")
    add_table(doc,
        ["Mês/2025", "Share EUA", "vs mesmo mês 2024"],
        [
            ["Janeiro", "34,9%", "-5,8 pp"],
            ["Março",   "41,7%", "-3,8 pp"],
            ["Junho",   "41,0%", "+2,8 pp"],
            ["Agosto",  "32,4%", "-11,3 pp"],
            ["Setembro","23,1%", "-19,1 pp"],
            ["Outubro", "21,8%", "-13,1 pp"],
            ["Novembro","18,5%", "-29,2 pp"],
            ["Dezembro","23,0%", "-18,6 pp"],
        ],
        col_widths=[4, 4, 5],
    )
    label(doc, "Visual sugerido")
    body(doc, "Gráfico de linha dupla (2024 vs 2025 mês a mês) com área sombreada da queda")
    label(doc, "Aba de referência")
    body(doc, "9_Mensal_EUA_vs_Outros")

    # Slide 10
    divider(doc)
    slide_title(doc, 10, "Quantificando a perda: análise contrafactual",
                "Estimativa de perda: US$ 124 milhões no segundo semestre de 2025")
    label(doc, "Metodologia (Event Study / Contrafactual)")
    bullet(doc, "Baseline: share dos EUA no H2/2024 = 40,7%")
    bullet(doc, "Total exportado florestal SC no H2/2025 = US$ 869 mi")
    bullet(doc, "Exportação esperada para EUA (se baseline mantido): US$ 354 mi")
    bullet(doc, "Exportação real para EUA no H2/2025: US$ 230 mi")
    bullet(doc, "Diferença (perda estimada): US$ 124 mi", level=1)
    label(doc, "Produtos mais afetados")
    add_table(doc,
        ["Produto", "Queda EUA 2024→2025"],
        [
            ["Carpintaria para construção", "-US$ 80 mi (-28%)"],
            ["Madeira em forma",            "-US$ 45 mi (-36%)"],
            ["Madeira compensada",          "-US$ 15 mi (-16%)"],
            ["Móveis",                      "-US$ 14 mi (-12%)"],
        ],
        col_widths=[9, 6],
    )
    label(doc, "Aba de referência")
    body(doc, "10_Queda_Produtos_EUA")

    # Slide 11
    divider(doc)
    slide_title(doc, 11, "Desvio de comércio: quem ganhou o que os EUA perderam",
                "Parte do comércio foi desviada — mas só parte")
    body(doc, "De US$ 124 mi perdidos nos EUA no H2/2025, outros mercados absorveram ~US$ 70 mi:")
    add_table(doc,
        ["País (ganhou share)", "Crescimento H1→H2 2025"],
        [
            ["México",        "+US$ 22 mi (+44%)"],
            ["África do Sul", "+US$ 9 mi (+150%)"],
            ["Índia",         "+US$ 8 mi (+76%)"],
            ["Espanha",       "+US$ 7 mi (+51%)"],
            ["Emirados",      "+US$ 7 mi (+44%)"],
            ["Paraguai",      "+US$ 6 mi (+28%)"],
            ["Marrocos",      "+US$ 4 mi (+95%)"],
        ],
        col_widths=[8, 7],
    )
    label(doc, "Saldo do desvio")
    bullet(doc, "Recuperado via outros mercados: ~US$ 70 mi")
    bullet(doc, "Destruição de comércio (não realizado): ~US$ 54 mi")
    label(doc, "Mensagem")
    quote_box(doc, "O mercado respondeu. Mas não foi suficiente para compensar. "
                   "Diversificação planejada com antecedência teria mitigado a perda.")
    label(doc, "Aba de referência")
    body(doc, "11_Desvio_Comercio")

    # Slide 12
    divider(doc)
    slide_title(doc, 12, "Concentração histórica nos EUA: uma vulnerabilidade conhecida",
                "A dependência dos EUA vinha crescendo desde 2015 — até 2025 forçar a mudança")
    add_table(doc,
        ["Ano", "Share EUA nas exp. florestais SC"],
        [
            ["2015", "36,4%"], ["2017", "39,0%"], ["2019", "42,7%"],
            ["2020", "47,4%"], ["2021", "50,0% (pico)"],
            ["2023", "40,9%"], ["2024", "40,9%"], ["2025", "33,3%"],
        ],
        col_widths=[4, 8],
    )
    label(doc, "Fala do orador")
    quote_box(doc, "Em 2021, metade de tudo que SC exportava em florestal ia para um único país. "
                   "Isso nunca foi uma estratégia — foi uma consequência do boom da construção civil americana. "
                   "As tarifas de 2025 forçaram uma correção que o mercado deveria ter feito antes.")
    label(doc, "Aba de referência")
    body(doc, "8_Concentracao_EUA")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # BLOCO 4
    # ══════════════════════════════════════════════════════════════════════════
    bloco(doc, "BLOCO 4 — Agenda e Conexão com Inovação (tema SEAFLOR)")

    # Slide 13
    divider(doc)
    slide_title(doc, 13, "O que a inovação já entregou para estes números",
                "Por trás dos US$ 1,86 bi estão décadas de P&D florestal")
    add_table(doc,
        ["Inovação", "Impacto nos números"],
        [
            ["Melhoramento genético (IMA Pinus/Eucalyptus)", "Madeira de maior qualidade → preço médio mais alto"],
            ["Manejo de precisão (sensoriamento remoto, drones)", "Redução de custo → margem para competir em preços pressionados por tarifa"],
            ["MDF e painéis (transformação industrial)", "Diversificação de produtos além da tora"],
            ["Certificações FSC/PEFC", "Acesso a mercados europeus e americanos exigentes"],
            ["CLT e construção em madeira", "Próxima onda de exportação — mercado crescente"],
        ],
        col_widths=[8, 8],
    )

    # Slide 14
    divider(doc)
    slide_title(doc, 14, "Agenda estratégica: o que os dados sugerem",
                "Três prioridades para o complexo florestal SC")
    label(doc, "1. Diversificação ativa de destinos")
    bullet(doc, "Base Florestal: Gini 0,94 — dois países concentram 90% das exportações")
    bullet(doc, "Oportunidades: Índia (+76%), Marrocos (+95%), África do Sul (+150%) no H2/2025")
    bullet(doc, "Ação: missões comerciais e prospecção ativa nos mercados emergentes")
    label(doc, "2. Adensamento tecnológico da cadeia")
    bullet(doc, "Madeira serrada (US$ 368 mi) tem menor valor agregado que carpintaria (US$ 254 mi)")
    bullet(doc, "Cada degrau na cadeia (tora → serrada → compensado → produto acabado) multiplica valor")
    bullet(doc, "Ação: incentivos para processamento local antes da exportação")
    label(doc, "3. Monitoramento do risco tarifário")
    bullet(doc, "Tarifas de 10% sobre produtos brasileiros nos EUA seguem vigentes")
    bullet(doc, "Próxima rodada de negociações pode elevar ou reduzir")
    bullet(doc, "Ação: mapeamento produto × alíquota HTS para cenários")

    # Slide 15
    divider(doc)
    slide_title(doc, 15, "Fechamento",
                "SC tem a floresta. O desafio é ampliar para quem vende.")
    label(doc, "Três números para levar")
    add_table(doc,
        ["Indicador", "Valor", "Significado"],
        [
            ["Exportações 2025",        "US$ 1,86 bi",   "15% do total exportado por SC"],
            ["Gini médio dos segmentos","0,86 – 0,94",   "Todos com alta concentração de destinos"],
            ["Custo estimado das tarifas (H2/2025)", "US$ 124 mi", "Perda vs. baseline 2024"],
        ],
        col_widths=[6.5, 3.5, 6],
    )
    label(doc, "Chamada para ação")
    body(doc, "Inovação florestal sem inteligência de mercado deixa valor na mesa. "
              "O Observatório FIESC está disponível para apoiar empresas e entidades do setor "
              "com dados atualizados e análises de comércio exterior.")
    body(doc, "obs.fiesc.com.br", bold=True, color=AZUL_FIESC)

    doc.add_page_break()

    # ── Apêndice ──────────────────────────────────────────────────────────────
    heading(doc, "Apêndice — Metodologia e Fontes", level=1)

    heading(doc, "Dados", level=2)
    bullet(doc, "Fonte primária: ComexStat / MDIC — balanca.economia.gov.br")
    bullet(doc, "Cobertura: Exportações e importações de SC por NCM, 2015–2025 (ano completo)")
    bullet(doc, "Categorização setorial: Dicionário SC Competitiva — Observatório FIESC")
    bullet(doc, "Processamento: Pipeline local Bronze → Gold → Datamarts (DuckDB + Python)")

    heading(doc, "Gini de concentração de destinos", level=2)
    bullet(doc, "Unidade de análise: Valor FOB exportado por país comprador, por ano e segmento")
    bullet(doc, "Fórmula: G = (2 × Σ(i × yᵢ)) / (n × Σyᵢ) − (n+1)/n  |  yᵢ ordenado crescente")
    bullet(doc, "HHI: Σ(sᵢ²) × 10.000  |  sᵢ = share de cada país")
    bullet(doc, "Interpretação: G → 0 = diversificado; G → 1 = concentrado")

    heading(doc, "Análise de impacto tarifário (Event Study / Contrafactual)", level=2)
    bullet(doc, "Baseline: share médio dos EUA no H2/2024 = 40,7%")
    bullet(doc, "Contrafactual: aplicação do baseline sobre o volume total exportado no H2/2025")
    bullet(doc, "Limitações: não controla por variações de câmbio, demanda do mercado americano "
                "e outros fatores simultâneos — estimativa de ordem de grandeza, não causal estrita")

    heading(doc, "Desvio de comércio", level=2)
    bullet(doc, "Método: comparação simples H1/2025 vs H2/2025 por país de destino")
    bullet(doc, "Pressuposto: H1/2025 ≈ período pré-efeito pleno das tarifas "
                "(anúncio em abril, impacto nas remessas a partir de julho/agosto)")

    heading(doc, "Arquivo de dados", level=2)
    body(doc, "SEAFLOR_2026_Comex_SC_Florestal.xlsx — 15 abas")
    body(doc, "Roteiro gerado em: junho 2026 | Observatório FIESC", italic=True, color=CINZA_TEXTO)

    # ── Salva ─────────────────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"Gerado: {OUT}")
    print(f"Tamanho: {OUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
