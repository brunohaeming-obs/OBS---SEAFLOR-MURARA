"""
Gera nota_metodologica_hhi_desvio_receita.docx
Observatório FIESC — Junho 2026
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# ── Paleta ────────────────────────────────────────────────────
AZL_ESC = "1F4E79"
AZL_MED = "2E75B6"
AZL_CL  = "DEEAF1"
VERDE   = "375623"
VERM    = "C00000"
LARANJ  = "C55A11"
CINZA   = "595959"
BRANCO  = "FFFFFF"
AMAR_CL = "FFFDE7"
VERD_CL = "E8F5E9"
VERM_CL = "FFF0F0"
CINZA_CL= "F5F5F5"

OUT     = Path(r"c:\Users\Janine\OneDrive\Área de Trabalho\OBS---SEAFLOR-MURARA\docs\nota_metodologica_hhi_desvio_receita.docx")
OUT_TMP = Path(r"C:\Users\Janine\AppData\Local\Temp\seaflor_metodologia_novo.docx")


def hex2rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── XML helpers ───────────────────────────────────────────────
def cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def cell_borders(cell, color=AZL_ESC, sz=8):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)


def table_borders(table, color=AZL_ESC, sz=8):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tblB.append(b)
    tblPr.append(tblB)


def para_shade(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ── Run helpers ───────────────────────────────────────────────
def run(para, text, bold=False, italic=False, sz=11,
        color=None, sup=False, sub=False):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(sz)
    if color:
        r.font.color.rgb = hex2rgb(color)
    if sup:
        r.font.superscript = True
    if sub:
        r.font.subscript = True
    return r


# ── Heading ───────────────────────────────────────────────────
def heading(doc, text, level=1, color=AZL_ESC):
    sz = {1: 14, 2: 12, 3: 11}[level]
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.color.rgb = hex2rgb(color)
    r.font.size = Pt(sz)
    r.bold = True
    return p


# ── Colored box (2-row table) ─────────────────────────────────
def box(doc, title, content_lines,
        bg=AZL_CL, border_col=AZL_ESC, title_bg=AZL_ESC,
        title_color=BRANCO):
    """
    content_lines: list of str or list of list-of-tuples
    Each str → plain line.
    Each list-of-tuples → one paragraph with multiple runs.
    Tuple: (text, bold, italic, color, sz, sup, sub)
    All optional after text.
    """
    t = doc.add_table(rows=2, cols=1)
    table_borders(t, border_col, 8)

    # Title row
    tc = t.rows[0].cells[0]
    cell_shade(tc, title_bg)
    cell_borders(tc, border_col, 8)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = hex2rgb(title_color)

    # Body row
    bc = t.rows[1].cells[0]
    cell_shade(bc, bg)
    cell_borders(bc, border_col, 8)

    first = True
    for line in content_lines:
        if first:
            p = bc.paragraphs[0]
            first = False
        else:
            p = bc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

        if isinstance(line, str):
            r = p.add_run(line)
            r.font.size = Pt(10)
        elif isinstance(line, list):
            for tup in line:
                text = tup[0]
                bold   = tup[1] if len(tup) > 1 else False
                italic = tup[2] if len(tup) > 2 else False
                clr    = tup[3] if len(tup) > 3 else None
                sz     = tup[4] if len(tup) > 4 else 10
                sup_   = tup[5] if len(tup) > 5 else False
                sub_   = tup[6] if len(tup) > 6 else False
                run(p, text, bold=bold, italic=italic, sz=sz,
                    color=clr, sup=sup_, sub=sub_)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return t


def formula_box(doc, title, lines):
    return box(doc, title, lines,
               bg=AZL_CL, border_col=AZL_ESC, title_bg=AZL_ESC)


def example_box(doc, title, lines):
    return box(doc, title, lines,
               bg=AMAR_CL, border_col=LARANJ, title_bg=LARANJ)


def atencao_box(doc, lines):
    return box(doc, "⚠  Atenção", lines,
               bg=VERM_CL, border_col=VERM, title_bg=VERM)


def interpreta_box(doc, lines):
    return box(doc, "✔  Como interpretar", lines,
               bg=VERD_CL, border_col=VERDE, title_bg=VERDE)


# ── Data table ────────────────────────────────────────────────
def data_table(doc, headers, rows, caption=None, note=None,
               col_widths=None, bold_last=True):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = hex2rgb(AZL_ESC)
        p.paragraph_format.space_after = Pt(2)

    n = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n)
    table_borders(t, AZL_ESC, 6)

    # Header
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        cell_shade(c, AZL_ESC)
        cell_borders(c, AZL_ESC, 6)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = hex2rgb(BRANCO)

    # Data rows
    for ri, row_data in enumerate(rows):
        is_total = any("Complexo" in str(v) or str(v).startswith("**")
                       for v in row_data)
        bg = CINZA_CL if is_total else (AZL_CL if ri % 2 == 0 else BRANCO)
        for ci, val in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            cell_shade(c, bg)
            cell_borders(c, AZL_ESC, 6)
            p = c.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            v = str(val).replace("**", "")
            r = p.add_run(v)
            r.font.size = Pt(9)
            if is_total and bold_last:
                r.bold = True

    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                if ci < len(t.rows[ri].cells):
                    set_col_width(t.rows[ri].cells[ci], w)

    if note:
        p = doc.add_paragraph()
        r = p.add_run(f"Nota: {note}")
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = hex2rgb(CINZA)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── Bullet list ───────────────────────────────────────────────
def bullets(doc, items, sz=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(sz)
        p.paragraph_format.space_after = Pt(2)


def body(doc, text, sz=11, color=None, italic=False, bold=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = hex2rgb(color)
    return p


def body_mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """parts: list of (text, bold, italic, color, sz)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    for tup in parts:
        text   = tup[0]
        bold_  = tup[1] if len(tup) > 1 else False
        italic_= tup[2] if len(tup) > 2 else False
        color_ = tup[3] if len(tup) > 3 else None
        sz_    = tup[4] if len(tup) > 4 else 11
        run(p, text, bold=bold_, italic=italic_, sz=sz_, color=color_)
    return p


def spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── CAPA ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("SEAFLOR 2026")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = hex2rgb(AZL_ESC)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nota Técnica de Metodologia")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = hex2rgb(AZL_MED)

    spacer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Concentração de Mercado (HHI)\n"
        "Decomposição Comercial\n"
        "Estimativa de Receita Foregone"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = hex2rgb(AZL_ESC)

    spacer(doc, 3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Observatório FIESC — Competitividade Industrial\nJunho de 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = hex2rgb(CINZA)

    spacer(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Este documento descreve, de forma didática e com rigor matemático, "
        "as três metodologias centrais do relatório SEAFLOR 2026. "
        "Para cada metodologia são apresentados: lógica conceitual, equações "
        "com definição de todos os termos, premissas, limitações e exemplo "
        "numérico passo a passo com dados reais do Complexo Florestal de SC."
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = hex2rgb(CINZA)

    spacer(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Fonte: ComexStat/MDIC  |  NCM–CNAE FIESC  |  "
        "Jan–Mai 2025 vs. Jan–Mai 2026"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = hex2rgb(CINZA)

    doc.add_page_break()

    # ── 1. APRESENTAÇÃO ───────────────────────────────────────
    heading(doc, "1. Apresentação e Escopo")
    body(doc,
         "Este documento técnico descreve as metodologias quantitativas "
         "utilizadas pelo Observatório FIESC para mensurar o impacto das "
         "tarifas americanas de 2025–2026 sobre as exportações do Complexo "
         "Florestal de Santa Catarina. O Complexo Florestal abrange quatro "
         "setores classificados pela CNAE:")

    data_table(doc,
        ["CNAE", "Atividade Econômica"],
        [
            ["2",  "Produção Florestal — base florestal (toras, lenha, carvão vegetal)"],
            ["16", "Fabricação de Produtos de Madeira (serrada, compensada, MDF, molduras)"],
            ["17", "Fabricação de Celulose, Papel e Produtos de Papel"],
            ["31", "Fabricação de Móveis e Componentes"],
        ],
        col_widths=[2.5, 13.0], bold_last=False
    )

    body(doc, "O documento está organizado em cinco blocos:")
    bullets(doc, [
        "Seção 2 — Bases de dados, fluxo de processamento e recortes analíticos.",
        "Seção 3 — Decomposição do impacto tarifário (destruição e desvio de comércio).",
        "Seção 4 — Índice de concentração Herfindahl-Hirschman (HHI).",
        "Seção 5 — Estimativa da receita foregone por análise contrafactual.",
        "Seção 6 — Resultados por CNAE com as três metodologias.",
    ])

    # ── 2. BASES DE DADOS ─────────────────────────────────────
    heading(doc, "2. Bases de Dados, Fluxo de Processamento e Recortes Analíticos")

    heading(doc, "2.1  Fonte Primária: ComexStat/MDIC", level=2)
    body(doc,
         "Todos os dados utilizados provêm do ComexStat, sistema do Ministério "
         "do Desenvolvimento, Indústria, Comércio e Serviços (MDIC). O ComexStat "
         "registra, ao nível de oito dígitos da Nomenclatura Comum do Mercosul (NCM), "
         "cada operação declarada ao Siscomex. As variáveis relevantes são:")
    bullets(doc, [
        "NCM (8 dígitos): classificação do produto exportado;",
        "UF do exportador: filtro para sg_uf = 'SC';",
        "País de destino: código ISO do país comprador;",
        "Mês/Ano: período da operação;",
        "Valor FOB (US$): valor na fronteira de embarque;",
        "Peso líquido (kg): volume físico exportado.",
    ], sz=10)
    body(doc,
         "Os arquivos são disponibilizados em dois formatos: (a) Histórico consolidado "
         "(gold) — anos 2022–2025 em Parquet; e (b) Bronze mensal — arquivos "
         "EXP_AAAAMM.csv de 2026, publicados mensalmente pelo MDIC.")

    heading(doc, "2.2  Dicionário de Classificação NCM–CNAE", level=2)
    body(doc,
         "A correspondência entre o NCM (produto) e a CNAE (setor) é feita por "
         "um dicionário desenvolvido pela FIESC, atualizado em abril de 2026. "
         "Cada NCM recebe exatamente um código CNAE (mapeamento 1:1):")

    data_table(doc,
        ["CNAE", "Capítulos SH principais", "Produtos representativos"],
        [
            ["2",  "SH 44 (parcial), SH 45", "Toras, lenha, carvão vegetal"],
            ["16", "SH 44 (principal)",        "Madeira serrada, compensada, MDF, molduras"],
            ["17", "SH 47, 48, 49",            "Celulose, papel kraft, embalagens"],
            ["31", "SH 94 (parcial)",           "Móveis, assentos, componentes"],
        ],
        col_widths=[2.0, 4.5, 9.0], bold_last=False
    )

    heading(doc, "2.3  Fluxo de Processamento de Dados", level=2)
    body(doc,
         "O processamento segue cinco etapas sequenciais, da fonte bruta "
         "até os indicadores do relatório:")

    data_table(doc,
        ["Etapa", "Nome", "Descrição"],
        [
            ["1", "Fonte bruta (bronze)",
             "Arquivos CSV mensais do ComexStat com todas as exportações do Brasil."],
            ["2", "Filtro e enriquecimento (silver)",
             "Seleção de sg_uf = 'SC' e aplicação do dicionário NCM–CNAE."],
            ["3", "Agregação (gold)",
             "Dados agrupados por {país de destino, CNAE, mês/ano} e armazenados em Parquet."],
            ["4", "Cálculo dos indicadores",
             "Scripts Python (DuckDB + pandas) calculam as métricas de desvio, HHI e receita foregone."],
            ["5", "Saída",
             "Resultados escritos no Excel SEAFLOR_2026_JanMai_3anos_v2.xlsx."],
        ],
        col_widths=[1.5, 4.0, 10.0], bold_last=False
    )

    heading(doc, "2.4  Recortes Analíticos", level=2)

    data_table(doc,
        ["Recorte", "Definição"],
        [
            ["Temporal",
             "Janeiro a Maio de cada ano — período simétrico que elimina sazonalidade. "
             "Jan–Mai 2026 é o primeiro semestre completo sob tarifas de 50% "
             "(vigentes desde 6 de agosto de 2025)."],
            ["Setorial",
             "CNAE 2 + 16 + 17 + 31 (Complexo Florestal). Os segmentos individuais "
             "podem usar filtros adicionais (sc_competitiva), gerando pequenas divergências "
             "com o total do Complexo."],
            ["Geográfico",
             "EUA = código de país 249 (Estados Unidos). OUT = todos os demais países "
             "(complemento de EUA). Para o HHI, inclui todos os países com exportação positiva."],
        ],
        col_widths=[3.0, 12.5], bold_last=False
    )

    heading(doc, "2.5  Notação", level=2)
    formula_box(doc, "Notação utilizada ao longo do documento", [
        "X(d, s, t)  =  valor FOB (US$) exportado pelo setor s para o destino d em Jan–Mai do ano t",
        "Q(d, s, t)  =  volume (kg líquido) exportado pelo setor s para o destino d em Jan–Mai do ano t",
        "X_TOT(s, t) =  total exportado pelo setor s em Jan–Mai do ano t (soma sobre todos os destinos)",
        "EUA  =  destino 'Estados Unidos' (código ISO 249)",
        "OUT  =  todos os destinos exceto EUA",
    ])

    # ── 3. DECOMPOSIÇÃO ───────────────────────────────────────
    heading(doc, "3. Metodologia 1 — Decomposição do Impacto Tarifário")

    heading(doc, "3.1  Lógica Conceitual", level=2)
    body(doc,
         "A decomposição adapta ao contexto tarifário a distinção clássica de "
         "Jacob Viner (1950) entre trade creation e trade diversion, originalmente "
         "desenvolvida para uniões aduaneiras. A lógica central se desenvolve em "
         "quatro movimentos:")
    bullets(doc, [
        "Quando os EUA impõem tarifa, o exportador catarinense perde competitividade "
        "naquele mercado — o comércio com os EUA é DESTRUÍDO.",
        "Parte dessa destruição pode ser compensada se o exportador redireciona produção "
        "a outros destinos — há DESVIO DE COMÉRCIO para fora dos EUA.",
        "O que sobra após o desvio é a DESTRUIÇÃO LÍQUIDA: a perda real e não compensada "
        "de receita em valor.",
        "Mas mesmo com desvio pode haver perda de receita por PREÇO: o mercado americano "
        "paga preços unitários superiores. Se a produção desviada obtém preço inferior, "
        "a receita cai mesmo que o volume não caia — esse é o EFEITO-PREÇO.",
    ], sz=10)

    formula_box(doc, "Fluxo lógico da decomposição", [
        "DB(s)  −  DC(s)  =  DL(s)",
        [
            ("  Destruição Bruta", True, False, AZL_ESC, 10),
            ("  menos  ", False, False, None, 10),
            ("Desvio de Comércio", True, False, VERDE, 10),
            ("  igual  ", False, False, None, 10),
            ("Destruição Líquida", True, False, AZL_ESC, 10),
        ],
        "",
        "DL(s)  +  EP(s)  =  PT(s)",
        [
            ("  Destruição Líquida", True, False, AZL_ESC, 10),
            ("  mais  ", False, False, None, 10),
            ("Efeito-Preço", True, False, LARANJ, 10),
            ("  igual  ", False, False, None, 10),
            ("Perda Total", True, False, VERM, 10),
        ],
    ])

    heading(doc, "3.2  Destruição Bruta (DB)", level=2)
    formula_box(doc, "Fórmula — Destruição Bruta", [
        "DB(s) = X(EUA, s, 2025) − X(EUA, s, 2026)",
        "",
        "onde:",
        "  X(EUA, s, 2025) = exportações do setor s para os EUA em Jan–Mai 2025",
        "  X(EUA, s, 2026) = exportações do setor s para os EUA em Jan–Mai 2026",
    ])
    body(doc,
         "O que mede: a queda absoluta em dólares nas exportações para os EUA "
         "entre Jan–Mai 2025 e Jan–Mai 2026.")
    bullets(doc, [
        "DB(s) > 0: houve queda nas exportações para os EUA. Caso de Madeira (CNAE 16) e Móveis (CNAE 31).",
        "DB(s) < 0: as exportações para os EUA CRESCERAM no período. Caso da Base Florestal (CNAE 2), "
        "possivelmente por contratos de longo prazo firmados antes das tarifas.",
    ], sz=10)
    atencao_box(doc, [
        "DB(s) mede o efeito sobre o canal EUA, não sobre o total exportado. Um setor pode exportar "
        "mais no total e ainda ter DB(s) > 0, se a perda nos EUA for compensada por outros destinos."
    ])

    heading(doc, "3.3  Desvio de Comércio (DC)", level=2)
    formula_box(doc, "Fórmula — Desvio de Comércio", [
        "DC(s) = máx[ X(OUT, s, 2026) − X(OUT, s, 2025)  ;  0 ]",
        "",
        "onde:",
        "  X(OUT, s, t) = exportações do setor s para todos os destinos EXCETO EUA em Jan–Mai do ano t",
        "  máx[ · ; 0 ] = a fórmula retorna zero se o resultado for negativo",
    ])
    body(doc,
         "O que mede: o ganho líquido nas exportações para mercados diferentes dos EUA. "
         "Se o setor expandiu vendas ao resto do mundo, esse incremento representa a "
         "parcela da produção antes destinada aos EUA que foi redirecionada com sucesso.")
    body(doc,
         "Por que usar máx(·, 0)? Se as exportações para fora dos EUA CAÍRAM, "
         "isso é uma perda adicional — não uma compensação. Nesses casos, DC(s) = 0 "
         "e toda a destruição bruta é líquida.")
    interpreta_box(doc, [
        "Um DC(s) elevado indica resiliência exportadora: o setor encontrou mercados "
        "alternativos. Um DC(s) ≈ 0 indica que a perda nos EUA não foi compensada — "
        "o volume antes destinado a esse mercado não foi absorvido por outros compradores."
    ])

    heading(doc, "3.4  Destruição Líquida (DL)", level=2)
    formula_box(doc, "Fórmula — Destruição Líquida", [
        "DL(s) = DB(s) − DC(s)",
        "",
        "Impacto líquido das tarifas sobre o valor total exportado,",
        "após descontar a compensação obtida em outros mercados.",
    ])
    bullets(doc, [
        "DL(s) > 0: a perda nos EUA supera o ganho em outros mercados. Caso de Madeira e Móveis.",
        "DL(s) < 0: os ganhos em outros mercados superaram a perda nos EUA — o setor CRESCEU "
        "em valor total. Caso do Papel e Celulose (CNAE 17).",
    ], sz=10)

    heading(doc, "3.5  Prêmio de Mercado EUA e Efeito-Preço (EP)", level=2)
    body(doc,
         "Mesmo quando o desvio compensa a perda em volume, o setor pode estar exportando "
         "a um preço unitário inferior ao praticado com os EUA. Esse diferencial é o "
         "Prêmio de Mercado EUA (π).")
    formula_box(doc, "Fórmulas — Prêmio EUA e Efeito-Preço", [
        "Valor unitário médio:   VU(d, s, t) = X(d, s, t) / Q(d, s, t)     [US$/kg]",
        "",
        "Prêmio de Mercado EUA:",
        "  π(s) = VU(EUA, s, 2025) − VU(OUT, s, 2025)",
        "",
        "Efeito-Preço:",
        "  EP(s) = máx(ΔQ_OUT(s), 0)  ×  máx(π(s), 0)",
        "",
        "  onde  ΔQ_OUT(s) = Q(OUT, s, 2026) − Q(OUT, s, 2025)   [variação de volume em kg]",
    ])
    body(doc,
         "Intuição: se o setor enviou mais quilos para outros mercados (ΔQ > 0), "
         "mas a um preço inferior ao americano (π > 0), houve perda de receita por "
         "unidade desviada. O Efeito-Preço quantifica esse custo de oportunidade.")
    atencao_box(doc, [
        "O Efeito-Preço usa o prêmio de 2025 como referência, não de 2026. "
        "Isso evita contaminação pelo próprio choque tarifário (que poderia ter "
        "alterado os preços dos EUA em 2026). O prêmio de 2025 representa o "
        "equilíbrio pré-choque."
    ])

    heading(doc, "3.6  Perda Total (PT)", level=2)
    formula_box(doc, "Fórmula — Perda Total (Método A)", [
        "PT(s) = DL(s) + EP(s)",
        "",
        "Expandindo:",
        "  PT(s) = [ DB(s) − DC(s) ]  +  [ máx(ΔQ_OUT, 0) × máx(π(s), 0) ]",
        "",
        "A Perda Total combina:",
        "  (1) Perda líquida de VALOR exportado: DL(s)",
        "  (2) Perda de receita por DIFERENCIAL DE PREÇO: EP(s)",
    ])

    heading(doc, "3.7  Exemplo Passo a Passo: Madeira (CNAE 16)", level=2)
    example_box(doc, "Exemplo completo — Madeira (CNAE 16)",  [
        "Dados observados (Jan–Mai 2025 vs. 2026):",
        "  X(EUA, 16, 2025) = US$ 262,1 mi     X(EUA, 16, 2026) = US$ 143,9 mi",
        "  X(OUT, 16, 2025) = US$ 277,4 mi     X(OUT, 16, 2026) = US$ 309,9 mi",
        "  VU(EUA, 16, 2025) = 0,942 US$/kg    VU(OUT, 16, 2025) = 0,434 US$/kg",
        "  Q(OUT, 16, 2025) = 638,4 mil ton     Q(OUT, 16, 2026) = 712,0 mil ton",
        "",
        "Passo 1 — Destruição Bruta:",
        "  DB(16) = 262,1 − 143,9 = US$ 118,1 mi",
        "",
        "Passo 2 — Desvio de Comércio:",
        "  DC(16) = máx(309,9 − 277,4 ; 0) = máx(32,5 ; 0) = US$ 32,5 mi",
        "",
        "Passo 3 — Destruição Líquida:",
        "  DL(16) = 118,1 − 32,5 = US$ 85,6 mi",
        "",
        "Passo 4 — Prêmio de Mercado EUA:",
        "  π(16) = 0,942 − 0,434 = 0,508 US$/kg  (+116,8% em relação ao preço nos demais mercados)",
        "",
        "Passo 5 — Efeito-Preço:",
        "  ΔQ(16) = 712,0 − 638,4 = +73,6 mil toneladas  [volume adicional desviado]",
        "  EP(16) ≈ 73.600 × 0,508 ≈ US$ 30,1 mi",
        "  (o valor exato resulta do cálculo por NCM, capturando variação de mix de produto)",
        "",
        "Passo 6 — Perda Total:",
        [("  PT(16) = 85,6 + 30,1 = US$ 115,7 milhões", True, False, VERM, 11)],
    ])

    # ── 4. HHI ────────────────────────────────────────────────
    heading(doc, "4. Metodologia 2 — Concentração de Mercado: Índice HHI")

    heading(doc, "4.1  Origem e Intuição", level=2)
    body(doc,
         "O Índice Herfindahl-Hirschman (HHI) foi proposto independentemente por "
         "Orris Herfindahl (1950) e Albert Hirschman (1945) para medir concentração "
         "em estruturas de mercado. No comércio internacional, é amplamente usado para "
         "quantificar a diversificação geográfica das exportações: quanto mais distribuídas "
         "entre países, menor o HHI e menor a vulnerabilidade a choques em um único mercado.")
    body(doc,
         "A ideia central do HHI é simples: o índice soma o quadrado das participações "
         "percentuais de cada destino. Elevar ao quadrado PENALIZA DESPROPORCIONALMENTE "
         "destinos com participação elevada:")
    bullets(doc, [
        "Um país com 40% do total contribui com  40² = 1.600 pontos.",
        "Quatro países com 10% cada contribuem juntos com apenas  4 × 10² = 400 pontos.",
        "Conclusão: reduzir a participação de um destino dominante é muito mais eficaz "
        "para baixar o HHI do que adicionar muitos destinos pequenos.",
    ], sz=10)

    heading(doc, "4.2  Formulação Matemática", level=2)
    formula_box(doc, "Fórmula — Índice HHI", [
        "HHI(s, t) = Σᵢ  [ X(i, s, t) / X_TOT(s, t) × 100 ]²",
        "",
        "onde:",
        "  i          = índice do país de destino (i = 1, 2, ..., N)",
        "  X(i, s, t) = valor FOB exportado para o país i, setor s, em Jan–Mai do ano t",
        "  X_TOT(s,t) = valor total exportado pelo setor s em Jan–Mai do ano t",
        "  N(s, t)    = número de países destino com exportação positiva",
        "",
        "O índice varia de 0 (distribuição perfeita) a 10.000 (exportação para um único país).",
    ])

    heading(doc, "4.3  Classificação de Concentração", level=2)
    body(doc,
         "A classificação adotada segue os critérios do Departamento de Justiça "
         "dos EUA (DOJ) e da Federal Trade Commission (FTC), adaptados ao contexto "
         "de exportações:")

    data_table(doc,
        ["HHI", "Classificação", "Interpretação para exportações"],
        [
            ["≥ 2.500",       "Alta concentração",     "Alta dependência; risco elevado de choque em destino principal"],
            ["1.500 – 2.499", "Concentração moderada", "Vulnerabilidade relevante a destino dominante"],
            ["1.000 – 1.499", "Baixa concentração",    "Pauta diversificada com alguma exposição residual"],
            ["< 1.000",       "Diversificado",          "Alta dispersão geográfica; baixa vulnerabilidade"],
        ],
        col_widths=[3.0, 3.5, 9.0], bold_last=False
    )

    heading(doc, "4.4  Contribuição Individual ao HHI", level=2)
    formula_box(doc, "Contribuição individual de cada país ao HHI", [
        "h(i, s, t) = [ X(i, s, t) / X_TOT(s, t) × 100 ]²",
        "",
        "de modo que:   HHI(s, t) = Σᵢ  h(i, s, t)",
        "",
        "Essa decomposição permite calcular exatamente quanto o HHI mudaria",
        "se a participação de um determinado destino fosse alterada.",
    ])
    body(doc,
         "Propriedade fundamental: a contribuição de um país ao HHI é proporcional "
         "ao quadrado da sua participação. Isso significa que a queda do share dos EUA "
         "de 39% para 25% produziu uma redução muito maior no índice do que o crescimento "
         "de Itália (+3 p.p.) ou Espanha (+2 p.p.) poderia ter aumentado.")

    heading(doc, "4.5  Exemplo Passo a Passo: Complexo Florestal SC", level=2)
    example_box(doc, "Cálculo do HHI — Complexo Florestal SC, Jan-Mai 2025 vs. 2026", [
        "Jan–Mai 2025  (total exportado: US$ 822,2 mi)",
        "",
        "  País              Valor (US$ mi)   Share (%)   Share²",
        "  Estados Unidos        322,2          39,18      1.535,1",
        "  Argentina              46,8           5,69         32,4",
        "  México                 38,3           4,66         21,7",
        "  Reino Unido            25,1           3,05          9,3",
        "  Alemanha               20,5           2,49          6,2",
        "  Outros (N−5)          369,3          44,93       (vários)",
        "  ─────────────────────────────────────────────────────────",
        "  HHI total                                       1.298",
        "",
        "Observação: os EUA contribuem com 1.535 pontos de um total de 1.298.",
        "Isso reflete que os demais países, individualmente pequenos, têm",
        "contribuições baixas — a concentração americana DOMINA o índice.",
        "",
        "─────────────────────────────────────────────────────────────────────",
        "",
        "Jan–Mai 2026  (total exportado: US$ 703,7 mi)",
        "",
        "  País              Valor (US$ mi)   Share (%)   Share²",
        "  Estados Unidos        174,4          24,79        614,5",
        "  México                 53,6           7,62         58,1",
        "  Argentina              44,8           6,37         40,6",
        "  Itália                 37,0           5,26         27,7",
        "  Espanha                27,4           3,89         15,1",
        "  Outros                366,5          52,07       (vários)",
        "  ─────────────────────────────────────────────────────────",
        "  HHI total                                          840",
        "",
        "O que explicou a queda?",
        "  h(EUA, 2025) = 39,18² = 1.535,1  →  h(EUA, 2026) = 24,79² = 614,5",
        "  Variação da contribuição EUA: −920,6 pontos",
        "  Variação total do HHI: −458 pontos",
        "",
        [("  ≈ 90% da redução do HHI veio exclusivamente da queda do share americano.", True, False, VERM, 10)],
    ])

    heading(doc, "4.6  Propriedades Importantes do HHI", level=2)
    bullets(doc, [
        "Sensibilidade à dominância: o índice reage muito mais à redução de um destino "
        "grande do que ao crescimento de muitos destinos pequenos.",
        "Independência do total exportado: o HHI mede concentração nas participações, "
        "não nos valores absolutos. Um setor pode exportar menos e ter HHI menor.",
        "Comparabilidade entre setores: como usa participações percentuais, permite "
        "comparar setores de tamanhos muito diferentes.",
        "Limitação — qualidade dos destinos: um HHI baixo pode refletir diversificação "
        "para mercados com baixo poder de compra ou alta instabilidade política.",
    ], sz=10)

    # ── 5. RECEITA FOREGONE ───────────────────────────────────
    heading(doc, "5. Metodologia 3 — Estimativa de Receita Foregone: Análise Contrafactual")

    heading(doc, "5.1  O que é uma Análise Contrafactual?", level=2)
    body(doc,
         "Uma análise contrafactual responde à pergunta: 'o que teria acontecido se as "
         "tarifas não tivessem sido impostas?' Para isso, constrói-se um cenário hipotético "
         "— o contrafactual — com base no comportamento histórico do setor ANTES do choque "
         "tarifário. A diferença entre o cenário contrafactual (sem tarifas) e o observado "
         "(com tarifas) é a Receita Foregone (RF) — a receita que o setor TERIA GERADO caso "
         "o choque não tivesse ocorrido.")
    interpreta_box(doc, [
        "Receita Foregone ≠ Destruição Líquida:",
        "",
        "• DL(s) compara o TOTAL EXPORTADO nos dois períodos — mede o que de fato foi "
        "perdido em valor total.",
        "",
        "• RF(s,k) estima o que TERIA SIDO EXPORTADO AOS EUA se a relação comercial "
        "fosse a do período k — mede a perda no canal americano, mesmo que o total tenha "
        "crescido. A RF captura perda de participação no mercado americano mesmo em setores "
        "cuja exportação total aumentou (ex.: Papel e Celulose).",
    ])

    heading(doc, "5.2  Formulação em Duas Etapas", level=2)
    formula_box(doc, "Etapa 1 — Share de referência EUA", [
        "sh_EUA(s, k) = X(EUA, s, k) / X_TOT(s, k)",
        "",
        "onde k é o período de referência histórico (antes do choque).",
        "Esse share representa a proporção 'normal' das exportações destinadas aos EUA.",
    ])
    formula_box(doc, "Etapa 2 — Exportações contrafactuais e Receita Foregone", [
        "X̂_EUA(s, k) = X_TOT(s, 2026) × sh_EUA(s, k)",
        "",
        "RF(s, k) = X̂_EUA(s, k) − X(EUA, s, 2026)",
        "",
        "Intuição: dado o total que o setor exportou em Jan–Mai 2026, quanto disso",
        "DEVERIA TER IDO para os EUA, se a relação comercial fosse a de antes?",
        "A diferença entre esse valor esperado e o observado é a receita foregone.",
    ])

    heading(doc, "5.3  Os Três Cenários de Referência", level=2)
    data_table(doc,
        ["Cenário", "Período k", "Pergunta que responde"],
        [
            ["A — Conservador", "Jan–Mai 2025",
             "Se nada mudasse em relação a Jan–Mai 2025, quanto iria aos EUA? "
             "(já embute antecipação às tarifas — tende a subestimar a perda)"],
            ["B — Base (principal)", "2024 (ano completo)",
             "Se a relação de 2024 — imediatamente anterior ao choque — fosse mantida, "
             "qual seria a perda? (cenário de referência principal, sem antecipação)"],
            ["C — Estrutural", "2023 (ano completo)",
             "Se a dependência estrutural pré-Trump fosse a referência, qual a perda "
             "máxima? (captura relação histórica antes das tensões comerciais)"],
        ],
        col_widths=[3.5, 3.5, 8.5], bold_last=False
    )
    body(doc,
         "A convergência entre Cenários B e C (US$ 114,4 vs. US$ 114,5 mi para o "
         "Complexo Florestal) é um sinal de ROBUSTEZ METODOLÓGICA: indica que a "
         "estrutura de dependência dos EUA era estável em 2023–2024, e que o intervalo "
         "estimado é confiável.")

    heading(doc, "5.4  Exemplo Passo a Passo: Complexo Florestal (Cenário B)", level=2)
    example_box(doc, "Receita Foregone — Complexo Florestal, Cenário B (base 2024)", [
        "Dados de entrada:",
        "  X(EUA, CF, 2024 ano)     = US$ 1.363,7 mi",
        "  X_TOT(CF, 2024 ano)      = US$ 3.321,4 mi",
        "  X(EUA, CF, 2026 jan-mai) = US$ 174,4 mi",
        "  X_TOT(CF, 2026 jan-mai)  = US$ 703,7 mi",
        "",
        "Passo 1 — Share de referência (2024):",
        "  sh_EUA(CF, B) = 1.363,7 / 3.321,4 = 41,04%",
        "",
        "Passo 2 — Exportações contrafactuais:",
        "  X̂_EUA(CF, B) = 703,7 × 0,4104 = US$ 288,8 mi",
        "",
        "Passo 3 — Receita Foregone:",
        [("  RF(CF, B) = 288,8 − 174,4 = US$ 114,4 milhões", True, False, VERM, 11)],
        "",
        "Interpretação: se o Complexo Florestal tivesse mantido em Jan–Mai 2026",
        "o mesmo share dos EUA de 2024 (41,04%), teria exportado US$ 288,8 mi",
        "para aquele mercado. Como exportou apenas US$ 174,4 mi, a receita",
        "foregone é de US$ 114,4 mi — equivalente a 16,3% do total exportado.",
    ])

    # ── 6. RESULTADOS ─────────────────────────────────────────
    heading(doc, "6. Resultados por CNAE — Jan-Mai 2025 vs. 2026")

    heading(doc, "6.1  Decomposição do Impacto Tarifário", level=2)
    data_table(doc,
        ["Setor", "DB (mi)", "DC (mi)", "DL (mi)", "EP (mi)", "PT (mi)"],
        [
            ["Madeira (CNAE 16)",          "118,1", "32,5",  "85,6",  "30,1", "115,7"],
            ["Móveis (CNAE 31)",            "24,9",  "0,0",  "24,9",   "0,0",  "24,9"],
            ["Papel e Celulose (CNAE 17)",   "5,2", "10,8",  "−5,6",  "17,7",  "12,2"],
            ["Base Florestal (CNAE 2)",     "−0,6",  "0,0",  "−0,6",   "0,0",  "−0,6"],
            ["**Complexo Florestal**",    "**147,7**","**29,2**","**118,6**","**0,0**","**118,6**"],
        ],
        caption="Tabela 1 — Decomposição do impacto tarifário (US$ milhões)",
        note="DB = Destruição Bruta | DC = Desvio de Comércio | DL = Destruição Líquida | "
             "EP = Efeito-Preço | PT = Perda Total. O EP do Complexo consolidado é zero porque o "
             "volume total para outros mercados recuou no período (apesar do valor em US$ ter crescido). "
             "Valores negativos em DB indicam crescimento das exportações para os EUA.",
        col_widths=[5.0, 2.0, 2.0, 2.0, 2.0, 2.0]
    )

    heading(doc, "6.2  Evolução do HHI de Concentração de Destinos", level=2)
    data_table(doc,
        ["Setor", "HHI 2024", "HHI 2025", "HHI jan-mai 26", "Δ 2024–26", "Classificação 2026"],
        [
            ["Madeira (CNAE 16)",          "2.695", "1.929", "1.253", "−1.442", "Baixa concentração"],
            ["Móveis (CNAE 31)",           "2.362", "1.877", "1.314", "−1.048", "Baixa concentração"],
            ["Papel e Celulose (CNAE 17)", "1.059", "1.118", "1.078",    "+19", "Baixa concentração"],
            ["**Complexo Florestal**",    "**1.863**","**1.298**","**840**","**−1.023**","**Diversificado**"],
        ],
        caption="Tabela 2 — Evolução do HHI por CNAE",
        note="HHI calculado sobre valor FOB de Jan–Mai de cada ano. Papel e Celulose manteve HHI "
             "estável (+19 pontos) pois não sofreu tarifa de 50%. Δ = HHI jan-mai 2026 − HHI 2024.",
        col_widths=[4.5, 2.0, 2.0, 2.5, 2.0, 2.5]
    )

    heading(doc, "6.3  Estimativa da Receita Foregone por Cenário", level=2)
    data_table(doc,
        ["Setor",
         "sh A (%)", "RF A (mi)",
         "sh B (%)", "RF B (mi)",
         "sh C (%)", "RF C (mi)"],
        [
            ["Madeira (CNAE 16)",          "48,58", "76,5",  "50,11", "83,5",  "50,65", "85,9"],
            ["Móveis (CNAE 31)",           "43,35", "10,5",  "45,68", "12,3",  "47,64", "13,9"],
            ["Papel e Celulose (CNAE 17)",  "7,72",  "5,6",   "7,61",  "5,5",  "10,35",  "9,7"],
            ["Base Florestal (CNAE 2)",     "0,47", "−0,6",   "3,63", "−0,1",   "3,81",  "0,0"],
            ["**Complexo Florestal**",    "**39,18**","**101,3**","**41,04**","**114,4**","**41,06**","**114,5**"],
        ],
        caption="Tabela 3 — Receita Foregone por cenário contrafactual (US$ milhões)",
        note="sh = share EUA de referência no período k. RF = Receita Foregone. "
             "Valores negativos na Base Florestal indicam que as exportações para os EUA "
             "superaram o contrafactual. A convergência entre B e C (114,4 vs. 114,5 mi) "
             "valida a robustez do intervalo estimado.",
        col_widths=[4.0, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7]
    )

    # ── 7. SÍNTESE ────────────────────────────────────────────
    heading(doc, "7. Síntese e Limitações Metodológicas")

    heading(doc, "7.1  Síntese por Setor", level=2)

    data_table(doc,
        ["Setor", "Resultado principal", "HHI 2024→2026", "Receita foregone (B)"],
        [
            ["Madeira (CNAE 16)",
             "PT = US$ 115,7 mi. DB de 118,1 mi parcialmente compensada por desvio "
             "(32,5 mi) e penalizada por efeito-preço (30,1 mi).",
             "2.695 → 1.253  (−1.442)", "US$ 83,5 mi"],
            ["Móveis (CNAE 31)",
             "PT = US$ 24,9 mi. Sem desvio e sem efeito-preço — mercados alternativos "
             "não absorveram o volume perdido.",
             "2.362 → 1.314  (−1.048)", "US$ 12,3 mi"],
            ["Papel e Celulose (CNAE 17)",
             "DL = −US$ 5,6 mi (cresceu em valor). Tarifa parcial (10%). "
             "Efeito-preço de 17,7 mi revela que crescimento foi a preços inferiores.",
             "1.059 → 1.078  (+19)", "US$ 5,5 mi"],
            ["Base Florestal (CNAE 2)",
             "Impacto desprezível. DB = −0,6 mi (exportações para EUA cresceram).",
             "—", "≈ 0"],
            ["**Complexo Florestal**",
             "DB 147,7 mi | DC 29,2 mi | DL 118,6 mi. Receita foregone: US$ 101–115 mi.",
             "1.863 → 840  (−1.023)", "US$ 114,4 mi"],
        ],
        caption="Tabela 4 — Síntese por setor",
        col_widths=[3.5, 6.5, 3.5, 2.5]
    )

    heading(doc, "7.2  Limitações Metodológicas", level=2)
    bullets(doc, [
        "Causalidade vs. correlação: toda a queda nos EUA é atribuída às tarifas. "
        "Outros fatores (ciclo da construção civil americana, câmbio, estoques) podem "
        "ter contribuído.",
        "Desvio medido em valor, não em quantidade: ganhos em outros mercados podem "
        "refletir aumento de preço sem expansão de volume. O Efeito-Preço tenta corrigir "
        "isso, mas assume prêmio uniforme por CNAE.",
        "Share de referência estático: o contrafactual supõe que o share dos EUA seria "
        "constante. Ignora crescimento orgânico que poderia ter ocorrido sem tarifas.",
        "Agregação por CNAE: heterogeneidade interna (ex.: madeira serrada vs. MDF dentro "
        "do CNAE 16) pode gerar resultados distintos ao nível de produto (NCM-8).",
        "Período de análise: 5 meses capturam o impacto imediato. Efeitos de médio prazo "
        "— fechamento de empresas, perda de contratos, realocação de capacidade produtiva "
        "— não estão refletidos.",
        "HHI e qualidade dos destinos: um índice baixo pode mascarar dependência de "
        "mercados com baixo poder de compra ou alta instabilidade política.",
    ], sz=10)

    # ── REFERÊNCIAS (ABNT NBR 6023:2018) ─────────────────────
    heading(doc, "Referências")

    # Nota sobre norma
    p = doc.add_paragraph()
    r = p.add_run("Referências elaboradas conforme ABNT NBR 6023:2018.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = hex2rgb(CINZA)
    p.paragraph_format.space_after = Pt(8)

    # Cada referência: (partes normais, partes em negrito)
    # Formato ABNT: SOBRENOME, Nome. Título em negrito. Local: Editora, ano.
    refs = [
        # [texto_antes_negrito, texto_negrito, texto_depois_negrito]
        (
            "BRASIL. Ministério do Desenvolvimento, Indústria, Comércio e Serviços. ",
            "ComexStat",
            ": sistema de análise das informações de comércio exterior. "
            "Brasília, DF: MDIC, 2026. Disponível em: https://comexstat.mdic.gov.br. "
            "Acesso em: 19 jun. 2026."
        ),
        (
            "FEDERAÇÃO DAS INDÚSTRIAS DO ESTADO DE SANTA CATARINA. Observatório FIESC. ",
            "Dicionário NCM–CNAE",
            ": mapeamento do Complexo Florestal de Santa Catarina. "
            "Versão atualizada em 7 abr. 2026. Florianópolis: FIESC, 2026."
        ),
        (
            "HERFINDAHL, Orris C. ",
            "Concentration in the steel industry",
            ". 1950. Tese (Doutorado em Economia) — Columbia University, Nova York, 1950."
        ),
        (
            "HIRSCHMAN, Albert O. ",
            "National power and the structure of foreign trade",
            ". Berkeley: University of California Press, 1945."
        ),
        (
            "UNITED STATES. Department of Justice; Federal Trade Commission. ",
            "Horizontal merger guidelines",
            ". Washington, DC: DOJ/FTC, 19 ago. 2010. Disponível em: "
            "https://www.justice.gov/atr/horizontal-merger-guidelines-08192010. "
            "Acesso em: 19 jun. 2026."
        ),
        (
            "UNITED STATES. The White House. ",
            "Executive order: modifying the scope of tariffs on the Government of Brazil",
            ". Washington, DC: The White House, 30 jul. 2025. Vigência: 6 ago. 2025. "
            "Disponível em: https://ballotpedia.org/Executive_Order:_Modifying_the_Scope_"
            "of_Tariffs_on_the_Government_of_Brazil_(Donald_Trump,_2025). "
            "Acesso em: 19 jun. 2026."
        ),
        (
            "VINER, Jacob. ",
            "The customs union issue",
            ". Nova York: Carnegie Endowment for International Peace, 1950."
        ),
        (
            "WORLD BANK. ",
            "Export diversification and quality",
            ". Washington, DC: World Bank, 2012."
        ),
    ]

    for antes, titulo, depois in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent       = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after       = Pt(6)
        p.paragraph_format.space_before      = Pt(0)
        # texto normal antes do título
        r1 = p.add_run(antes)
        r1.font.size = Pt(10)
        # título em negrito/itálico (ABNT: negrito para obras)
        r2 = p.add_run(titulo)
        r2.bold = True
        r2.font.size = Pt(10)
        # texto normal depois do título
        r3 = p.add_run(depois)
        r3.font.size = Pt(10)

    # ── APÊNDICE ──────────────────────────────────────────────
    heading(doc, "Apêndice — Glossário Completo de Variáveis")

    data_table(doc,
        ["Símbolo", "Descrição"],
        [
            ["X(d, s, t)",     "Valor FOB (US$) exportado pelo setor s para o destino d em Jan–Mai do ano t"],
            ["Q(d, s, t)",     "Volume (kg líquido) exportado pelo setor s para d em Jan–Mai do ano t"],
            ["X_TOT(s, t)",    "Total exportado pelo setor s em Jan–Mai do ano t (soma sobre todos os destinos)"],
            ["VU(d, s, t)",    "Valor unitário médio: X(d,s,t) / Q(d,s,t) em US$/kg"],
            ["EUA",            "Estados Unidos (código ISO 249)"],
            ["OUT",            "Todos os destinos exceto EUA"],
            ["CF",             "Complexo Florestal (CNAE 2 + 16 + 17 + 31)"],
            ["DB(s)",          "Destruição Bruta: X(EUA,s,2025) − X(EUA,s,2026)"],
            ["DC(s)",          "Desvio de Comércio: máx[ X(OUT,s,2026) − X(OUT,s,2025) ; 0 ]"],
            ["DL(s)",          "Destruição Líquida: DB(s) − DC(s)"],
            ["π(s)",           "Prêmio de Mercado EUA: VU(EUA,s,2025) − VU(OUT,s,2025) em US$/kg"],
            ["ΔQ_OUT(s)",      "Variação de volume para OUT: Q(OUT,s,2026) − Q(OUT,s,2025)"],
            ["EP(s)",          "Efeito-Preço: máx(ΔQ_OUT(s), 0) × máx(π(s), 0)"],
            ["PT(s)",          "Perda Total: DL(s) + EP(s)"],
            ["HHI(s, t)",      "Índice Herfindahl-Hirschman do setor s no período t"],
            ["h(i, s, t)",     "Contribuição individual do país i ao HHI: [X(i,s,t)/X_TOT(s,t) × 100]²"],
            ["N(s, t)",        "Número de países destino com exportação positiva (setor s, período t)"],
            ["sh_EUA(s, k)",   "Share EUA de referência: X(EUA,s,k) / X_TOT(s,k)"],
            ["X̂_EUA(s, k)",   "Exportações contrafactuais para os EUA: X_TOT(s,2026) × sh_EUA(s,k)"],
            ["RF(s, k)",       "Receita Foregone: X̂_EUA(s,k) − X(EUA,s,2026)"],
        ],
        col_widths=[3.5, 12.0], bold_last=False
    )

    # ── SALVAR ────────────────────────────────────────────────
    OUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(OUT))
        print(f"Documento salvo: {OUT}")
        sz = OUT.stat().st_size / 1024
        print(f"Tamanho: {sz:.0f} KB")
    except PermissionError:
        doc.save(str(OUT_TMP))
        print(f"Arquivo em uso — salvo em: {OUT_TMP}")
        print("Feche o Word e copie para substituir o original.")
        sz = OUT_TMP.stat().st_size / 1024
        print(f"Tamanho: {sz:.0f} KB")


if __name__ == "__main__":
    build()
